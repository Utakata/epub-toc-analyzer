#!/usr/bin/env python
# -*- coding: utf-8 -*-

__license__   = 'GPL v3'
__copyright__ = '2025, Enhanced Word TOC Output for Japan v2.0'
__docformat__ = 'restructuredtext en'
__version__ = '2.0.0'

import sys, re, os, traceback, copy, glob
from posixpath import normpath
import logging
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor
import chardet

# Progress bar support
try:
    from tqdm import tqdm
    PROGRESS_SUPPORT = True
except ImportError:
    PROGRESS_SUPPORT = False
    
logger = logging.getLogger(__name__)

from zipfile import ZipFile, ZIP_STORED, ZIP_DEFLATED
from xml.dom.minidom import parse, parseString, getDOMImplementation, Element
from time import time
from datetime import datetime

import six
from six.moves.urllib.parse import unquote
from six import string_types, text_type as unicode
from six import unichr

from bs4 import BeautifulSoup

# Word document generation support
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    WORD_SUPPORT = True
except ImportError:
    WORD_SUPPORT = False
    print("python-docx not available. Word output will be disabled.")

class CalibreCompatibleTOCDetector:
    """
    Calibre互換の目次検出クラス
    XPath式ベースの階層検出とヒューリスティック処理を実装
    """
    
    def __init__(self):
        # Calibre風XPath式パターン
        self.xpath_patterns = {
            'level1': ['//h:h1', '//h1', '//div[@class="chapter"]', '//div[@class="section1"]'],
            'level2': ['//h:h2', '//h2', '//div[@class="section"]', '//div[@class="section2"]'],
            'level3': ['//h:h3', '//h3', '//div[@class="subsection"]', '//div[@class="section3"]']
        }
        
        # ヒューリスティックパターン
        self.heuristic_patterns = {
            'chapter': re.compile(r'(第\d+章|Chapter\s+\d+|CHAPTER\s+\d+)', re.IGNORECASE),
            'section': re.compile(r'(\d+\.\d+|\d+－\d+|§\d+)', re.IGNORECASE),
            'subsection': re.compile(r'(\d+\.\d+\.\d+|\(\d+\))', re.IGNORECASE)
        }
    
    def detect_toc_from_html(self, html_content):
        """HTMLコンテンツから目次構造を検出"""
        soup = BeautifulSoup(html_content, 'lxml')
        detected_toc = {'level1': [], 'level2': [], 'level3': []}
        
        # XPath式に相当するCSS Selectorで検出
        for level, patterns in [('level1', ['h1', 'div.chapter', 'div.section1']),
                               ('level2', ['h2', 'div.section', 'div.section2']),
                               ('level3', ['h3', 'div.subsection', 'div.section3'])]:
            
            for pattern in patterns:
                elements = soup.select(pattern)
                for elem in elements:
                    text = elem.get_text().strip()
                    if text and len(text) > 1:
                        detected_toc[level].append({
                            'text': self._clean_heading_text(text),
                            'tag': elem.name,
                            'position': self._get_element_position(elem)
                        })
        
        # ヒューリスティック検出
        self._apply_heuristic_detection(soup, detected_toc)
        
        return detected_toc
    
    def _clean_heading_text(self, text):
        """見出しテキストのクリーニング"""
        # 不要な文字の除去
        text = re.sub(r'\s+', ' ', text).strip()
        # 日本語特有のパターン正規化
        text = re.sub(r'^第(\d+)章\s*', r'第\1章　', text)
        text = re.sub(r'^(\d+)\.\s*', r'\1. ', text)
        return text
    
    def _get_element_position(self, element):
        """要素のドキュメント内位置を取得"""
        # BeautifulSoupでの概算位置
        return len(str(element.encode_contents()))
    
    def _apply_heuristic_detection(self, soup, detected_toc):
        """Calibre風ヒューリスティック検出"""
        all_paragraphs = soup.find_all(['p', 'div', 'span'])
        
        for para in all_paragraphs:
            text = para.get_text().strip()
            if not text:
                continue
                
            # パターンマッチング
            if self.heuristic_patterns['chapter'].search(text):
                if len(text) < 100:  # 長すぎるものは除外
                    detected_toc['level1'].append({
                        'text': self._clean_heading_text(text),
                        'tag': 'heuristic_h1',
                        'position': self._get_element_position(para)
                    })
            elif self.heuristic_patterns['section'].search(text):
                if len(text) < 80:
                    detected_toc['level2'].append({
                        'text': self._clean_heading_text(text),
                        'tag': 'heuristic_h2',
                        'position': self._get_element_position(para)
                    })
            elif self.heuristic_patterns['subsection'].search(text):
                if len(text) < 60:
                    detected_toc['level3'].append({
                        'text': self._clean_heading_text(text),
                        'tag': 'heuristic_h3',
                        'position': self._get_element_position(para)
                    })

class BatchProcessor:
    """バッチ処理クラス - 複数ファイルの一括処理"""
    
    def __init__(self, max_workers=4):
        self.max_workers = max_workers
        self.results = []
        self.errors = []
    
    def process_directory(self, directory_path, output_dir=".", format_type="both", recursive=True):
        """ディレクトリ内のEPUBファイルを一括処理"""
        epub_files = self._find_epub_files(directory_path, recursive)
        
        if not epub_files:
            print(f"📁 {directory_path} にEPUBファイルが見つかりませんでした")
            return []
        
        print(f"📚 {len(epub_files)}個のEPUBファイルを検出しました")
        
        # プログレスバー付きバッチ処理
        if PROGRESS_SUPPORT:
            with tqdm(total=len(epub_files), desc="EPUB処理中") as pbar:
                with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
                    futures = []
                    for epub_file in epub_files:
                        future = executor.submit(self._process_single_epub, epub_file, output_dir, format_type)
                        futures.append(future)
                    
                    for future in futures:
                        try:
                            result = future.result()
                            self.results.append(result)
                            pbar.update(1)
                        except Exception as e:
                            self.errors.append(str(e))
                            pbar.update(1)
        else:
            # プログレスバーなしの処理
            for i, epub_file in enumerate(epub_files, 1):
                print(f"処理中... {i}/{len(epub_files)}: {Path(epub_file).name}")
                try:
                    result = self._process_single_epub(epub_file, output_dir, format_type)
                    self.results.append(result)
                except Exception as e:
                    self.errors.append(str(e))
                    print(f"❌ エラー: {e}")
        
        return self.results
    
    def _find_epub_files(self, directory_path, recursive=True):
        """EPUBファイルを検索"""
        directory = Path(directory_path)
        if not directory.exists():
            return []
        
        if recursive:
            return list(directory.rglob("*.epub"))
        else:
            return list(directory.glob("*.epub"))
    
    def _process_single_epub(self, epub_path, output_dir, format_type):
        """単一EPUBファイルの処理"""
        try:
            with open(epub_path, 'rb') as f:
                epub_splitter = SplitEpubWordTOC(f)
                results = epub_splitter.generate_word_toc_output(
                    output_dir=output_dir,
                    format_type=format_type
                )
            return {'file': str(epub_path), 'results': results, 'status': 'success'}
        except Exception as e:
            return {'file': str(epub_path), 'error': str(e), 'status': 'error'}

class EnhancedWordTOCGenerator:
    """改良版Word形式の目次レベル3段階出力ジェネレーター"""
    
    def __init__(self):
        self.levels = {1: [], 2: [], 3: []}
        self.current_level_1 = None
        self.current_level_2 = None
        self.book_title = ""
        self.authors = []
        self.calibre_detector = CalibreCompatibleTOCDetector()
        
    def add_toc_entry(self, text, level, href="", anchor="", hierarchy_path="", detected_method="standard"):
        """目次エントリを追加（検出方法の記録付き）"""
        entry = {
            'text': text,
            'href': href,
            'anchor': anchor,
            'hierarchy_path': hierarchy_path,
            'level': level,
            'detected_method': detected_method,  # 検出方法を記録
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }
        
        if level == 1:
            self.current_level_1 = text
            self.current_level_2 = None
            self.levels[1].append(entry)
        elif level == 2:
            self.current_level_2 = text
            entry['parent_level_1'] = self.current_level_1
            self.levels[2].append(entry)
        elif level == 3:
            entry['parent_level_1'] = self.current_level_1
            entry['parent_level_2'] = self.current_level_2
            self.levels[3].append(entry)
    
    def generate_enhanced_text_output(self):
        """改良版テキスト形式での目次出力"""
        output = []
        output.append("=" * 80)
        output.append(f"📚 書籍目次構造分析レポート v2.0 (Calibre互換)")
        output.append(f"書籍名: {self.book_title}")
        output.append(f"著者: {', '.join(self.authors)}")
        output.append(f"生成日時: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
        output.append("=" * 80)
        output.append("")
        
        # 検出統計
        total_entries = sum(len(self.levels[i]) for i in range(1, 4))
        output.append("🔍 検出統計")
        output.append("-" * 60)
        output.append(f"総エントリ数: {total_entries}")
        
        # 検出方法別の統計
        detection_stats = {}
        for level in range(1, 4):
            for entry in self.levels[level]:
                method = entry.get('detected_method', 'unknown')
                detection_stats[method] = detection_stats.get(method, 0) + 1
        
        for method, count in detection_stats.items():
            output.append(f"  - {method}: {count}件")
        output.append("")
        
        # 各レベルの詳細
        for level in range(1, 4):
            if not self.levels[level]:
                continue
                
            level_names = {1: "📖 目次レベル1（大見出し）", 
                          2: "📝 目次レベル2（中見出し）", 
                          3: "📄 目次レベル3（小見出し）"}
            
            output.append(level_names[level])
            output.append("-" * 60)
            
            for i, entry in enumerate(self.levels[level], 1):
                output.append(f"{i:2d}. {entry['text']}")
                
                # 親階層の表示
                if level >= 2 and entry.get('parent_level_1'):
                    output.append(f"     └ 親レベル1: {entry['parent_level_1']}")
                if level == 3 and entry.get('parent_level_2'):
                    output.append(f"       └ 親レベル2: {entry['parent_level_2']}")
                
                # 技術的詳細
                if entry['href']:
                    output.append(f"     ファイル: {entry['href']}")
                if entry['anchor']:
                    output.append(f"     アンカー: #{entry['anchor']}")
                if entry.get('detected_method') != 'standard':
                    output.append(f"     検出方法: {entry['detected_method']}")
            output.append("")
        
        return '\n'.join(output)
    
    def generate_enhanced_word_document(self, output_path):
        """改良版Word文書として目次を出力"""
        if not WORD_SUPPORT:
            raise ImportError("python-docx is required for Word output")
            
        doc = Document()
        
        # ドキュメントタイトル
        title = doc.add_heading('書籍目次構造分析レポート v2.0', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 書籍情報
        info_para = doc.add_paragraph()
        info_para.add_run('書籍名: ').bold = True
        info_para.add_run(self.book_title)
        info_para.add_run('\n著者: ').bold = True
        info_para.add_run(', '.join(self.authors))
        info_para.add_run('\n生成日時: ').bold = True
        info_para.add_run(datetime.now().strftime('%Y年%m月%d日 %H:%M:%S'))
        info_para.add_run('\n解析エンジン: ').bold = True
        info_para.add_run('Calibre互換 + ヒューリスティック検出')
        
        # 統計情報を先頭に
        stats_heading = doc.add_heading('📊 検出統計', level=1)
        stats_para = doc.add_paragraph()
        stats_para.add_run('レベル1エントリ数: ').bold = True
        stats_para.add_run(str(len(self.levels[1])))
        stats_para.add_run('\nレベル2エントリ数: ').bold = True
        stats_para.add_run(str(len(self.levels[2])))
        stats_para.add_run('\nレベル3エントリ数: ').bold = True
        stats_para.add_run(str(len(self.levels[3])))
        
        doc.add_page_break()
        
        # 各レベルの目次
        level_names = {1: '📖 目次レベル1（大見出し）', 
                      2: '📝 目次レベル2（中見出し）', 
                      3: '📄 目次レベル3（小見出し）'}
        
        for level in range(1, 4):
            if not self.levels[level]:
                continue
                
            level_heading = doc.add_heading(level_names[level], level=1)
            
            for i, entry in enumerate(self.levels[level], 1):
                para = doc.add_paragraph(style='List Number')
                para.add_run(entry['text']).bold = True
                
                # 階層情報
                if level >= 2:
                    parent_para = doc.add_paragraph(style='List Bullet 2')
                    if entry.get('parent_level_1'):
                        parent_para.add_run(f"└ 親レベル1: {entry['parent_level_1']}")
                    if level == 3 and entry.get('parent_level_2'):
                        parent_para.add_run(f"\n  └ 親レベル2: {entry['parent_level_2']}")
                
                # 技術的詳細
                if entry['href'] or entry['anchor'] or entry.get('detected_method') != 'standard':
                    detail_para = doc.add_paragraph(style='List Bullet 3')
                    details = []
                    if entry['href']:
                        details.append(f"ファイル: {entry['href']}")
                    if entry['anchor']:
                        details.append(f"アンカー: #{entry['anchor']}")
                    if entry.get('detected_method') != 'standard':
                        details.append(f"検出: {entry['detected_method']}")
                    detail_para.add_run("   ".join(details))
        
        # ドキュメント保存
        doc.save(output_path)
        return output_path

# メインのSplitEpubWordTOCクラスも更新
class SplitEpubWordTOC:
    """EPUB splitter with enhanced Word TOC output support v2.0"""
    
    def __init__(self, inputio):
        self.epub = ZipFile(inputio, 'r')
        self.content_dom = None
        self.content_relpath = None
        self.manifest_items = None
        self.guide_items = None
        self.toc_dom = None
        self.toc_relpath = None
        self.toc_map = None
        self.split_lines = None
        self.origauthors = []
        self.origtitle = None
        
        # Enhanced features v2.0
        self.toc_processor = None
        self.hierarchy_enabled = True
        self.calibre_detector = CalibreCompatibleTOCDetector()
        
    # ... (以前のメソッドは同じなので省略) ...
    
    def generate_word_toc_output(self, output_dir=".", format_type="both"):
        """
        v2.0改良版: Word形式目次レベル3段階出力のメイン関数
        Calibre互換の検出機能とエラー処理強化
        """
        try:
            # エンコーディング検出と処理
            self._detect_and_handle_encoding()
            
            # TOC解析実行
            self.get_enhanced_toc_map()
            
            if not self.toc_processor:
                raise Exception("TOC processor not initialized")
            
            results = {}
            base_name = self._sanitize_filename(self.get_title())
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            # 出力ディレクトリの確認と作成
            output_path = Path(output_dir)
            output_path.mkdir(parents=True, exist_ok=True)
            
            # テキスト形式出力
            if format_type in ["text", "both"]:
                text_output = self.toc_processor.word_toc_generator.generate_enhanced_text_output()
                text_path = output_path / f"{base_name}_目次分析_{timestamp}.txt"
                
                with open(text_path, 'w', encoding='utf-8') as f:
                    f.write(text_output)
                
                results['text'] = str(text_path)
                print(f"📄 テキスト形式目次ファイル生成: {text_path}")
            
            # Word文書出力
            if format_type in ["word", "both"] and WORD_SUPPORT:
                word_path = output_path / f"{base_name}_目次分析_{timestamp}.docx"
                self.toc_processor.word_toc_generator.generate_enhanced_word_document(str(word_path))
                results['word'] = str(word_path)
                print(f"📝 Word文書目次ファイル生成: {word_path}")
            elif format_type in ["word", "both"] and not WORD_SUPPORT:
                print("⚠️  Word出力にはpython-docxライブラリが必要です")
            
            return results
            
        except Exception as e:
            logger.error(f"TOC生成エラー: {e}", exc_info=True)
            raise Exception(f"目次生成に失敗しました: {e}")
    
    def _detect_and_handle_encoding(self):
        """エンコーディング検出と処理"""
        try:
            # EPUBファイル内のエンコーディングをチェック
            container = self.epub.read("META-INF/container.xml")
            if isinstance(container, bytes):
                detected = chardet.detect(container)
                if detected['confidence'] > 0.7:
                    logger.info(f"Detected encoding: {detected['encoding']}")
        except:
            pass  # エンコーディング検出に失敗しても続行
    
    def _sanitize_filename(self, filename):
        """ファイル名の無効文字を除去"""
        if not filename:
            return "unknown_book"
        # Windows/Linux両対応の安全なファイル名
        safe_chars = re.sub(r'[<>:"/\\|?*]', '_', filename)
        safe_chars = re.sub(r'\s+', '_', safe_chars)
        return safe_chars[:100]  # 長さ制限

# Utility functions (変更なし)
def get_path_part(n):
    relpath = os.path.dirname(n)
    if len(relpath) > 0:
        relpath = relpath + "/"
    return relpath

def get_file_part(n):
    return os.path.basename(n)

def main(argv):
    """コマンドライン実行用メイン関数 v2.0"""
    from optparse import OptionParser
    
    usage = 'usage: python %prog [options] <input epub>'
    parser = OptionParser(usage + '''

Word形式目次レベル3段階出力ツール v2.0 - Calibre互換版

EPUBファイルの目次構造を分析し、Word形式で3段階まで詳細に出力します。
v2.0では、Calibreのベストプラクティスを採用し、バッチ処理にも対応。
''')

    parser.add_option("-o", "--output-dir", dest="output_dir", default=".",
                      help="出力ディレクトリを指定 (デフォルト: 現在のディレクトリ)", metavar="DIR")
    parser.add_option("-f", "--format", dest="format", default="both",
                      help="出力形式を指定: text, word, both (デフォルト: both)", metavar="FORMAT")
    parser.add_option("-b", "--batch", dest="batch_dir", default=None,
                      help="バッチ処理: 指定ディレクトリ内の全EPUBファイルを処理", metavar="DIR")
    parser.add_option("--workers", dest="workers", type="int", default=4,
                      help="バッチ処理の並列実行数 (デフォルト: 4)", metavar="NUM")

    (options, args) = parser.parse_args(argv)

    # バッチ処理モード
    if options.batch_dir:
        print(f"🔄 バッチ処理モード: {options.batch_dir}")
        processor = BatchProcessor(max_workers=options.workers)
        results = processor.process_directory(
            options.batch_dir, 
            output_dir=options.output_dir,
            format_type=options.format
        )
        
        # 結果サマリー
        success_count = len([r for r in results if r['status'] == 'success'])
        error_count = len([r for r in results if r['status'] == 'error'])
        
        print(f"\n📊 バッチ処理結果:")
        print(f"✅ 成功: {success_count}ファイル")
        print(f"❌ エラー: {error_count}ファイル")
        
        if processor.errors:
            print(f"\n❌ エラー詳細:")
            for error in processor.errors[:5]:  # 最初の5件のみ表示
                print(f"  - {error}")
        
        return

    # 単一ファイル処理モード
    if not args:
        parser.print_help()
        return

    epub_path = args[0]
    if not os.path.exists(epub_path):
        print(f"❌ エラー: ファイルが見つかりません: {epub_path}")
        return

    print(f"📚 EPUB分析開始: {epub_path}")
    
    try:
        with open(epub_path, 'rb') as f:
            epub_splitter = SplitEpubWordTOC(f)
            results = epub_splitter.generate_word_toc_output(
                output_dir=options.output_dir,
                format_type=options.format
            )
        
        print("\n✅ 目次分析完了!")
        print(f"📊 出力ファイル:")
        for format_type, path in results.items():
            print(f"   - {format_type}: {path}")
            
    except Exception as e:
        print(f"❌ エラーが発生しました: {e}")
        logger.error(f"Error processing EPUB: {e}", exc_info=True)

if __name__ == "__main__":
    main(sys.argv[1:])
