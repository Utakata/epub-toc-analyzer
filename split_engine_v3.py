#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
EPUB目次解析ツール v3.0 - 分割機能エンジン
動的目次解析・Wordファイル分割・マルチフォーマット出力対応
"""

__version__ = '3.0.0'
__license__ = 'GPL v3'
__copyright__ = '2025, Enhanced Split Engine v3.0'

import os
import re
import sys
import json
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple, Optional, Any
from dataclasses import dataclass
from collections import defaultdict

# Word処理
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.document import Document as DocxDocument
    from docx.text.paragraph import Paragraph
    import docx2txt
    WORD_SUPPORT = True
except ImportError:
    WORD_SUPPORT = False

# PDF生成
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph as RLParagraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

# EPUB生成
try:
    import ebooklib
    from ebooklib import epub
    EPUB_SUPPORT = True
except ImportError:
    EPUB_SUPPORT = False

# 既存v2.0モジュール
try:
    from epubsplit_word_toc_v2 import SplitEpubWordTOC, CalibreCompatibleTOCDetector
    V2_SUPPORT = True
except ImportError:
    V2_SUPPORT = False

@dataclass
class TOCEntry:
    """目次エントリデータクラス"""
    text: str
    level: int
    page_num: Optional[int] = None
    paragraph_index: Optional[int] = None
    style_name: Optional[str] = None
    parent_path: str = ""
    file_href: str = ""
    anchor: str = ""
    
    def __post_init__(self):
        if self.parent_path and self.text:
            self.full_path = f"{self.parent_path}/{self.text}"
        else:
            self.full_path = self.text

@dataclass 
class SplitConfig:
    """分割設定データクラス"""
    split_level: int = 1  # 分割する目次レベル
    output_format: str = "word"  # word, pdf, epub
    include_subsections: bool = True  # 下位レベルも含める
    preserve_formatting: bool = True  # 書式保持
    output_dir: str = "output"
    filename_pattern: str = "{index:02d}_{title}"
    max_filename_length: int = 50

class DynamicTOCAnalyzer:
    """動的目次構造解析クラス"""
    
    def __init__(self):
        self.toc_entries: List[TOCEntry] = []
        self.level_stats: Dict[int, int] = defaultdict(int)
        self.max_depth = 0
        self.recommended_split_level = 1
        
    def analyze_word_document(self, docx_path: str) -> Dict[str, Any]:
        """Wordファイルの目次構造を解析"""
        if not WORD_SUPPORT:
            raise ImportError("python-docx is required for Word analysis")
            
        document = Document(docx_path)
        self.toc_entries = []
        self.level_stats = defaultdict(int)
        
        # 見出しスタイルの検出
        heading_styles = self._detect_heading_styles(document)
        
        # パラグラフ解析
        for para_idx, paragraph in enumerate(document.paragraphs):
            if paragraph.style.name in heading_styles:
                level = heading_styles[paragraph.style.name]
                text = paragraph.text.strip()
                
                if text:  # 空でない見出しのみ
                    entry = TOCEntry(
                        text=text,
                        level=level,
                        paragraph_index=para_idx,
                        style_name=paragraph.style.name
                    )
                    self.toc_entries.append(entry)
                    self.level_stats[level] += 1
        
        # 階層構造の構築
        self._build_hierarchy()
        
        # 分析結果をまとめる
        return self._generate_analysis_report()
    
    def analyze_epub_toc(self, epub_path: str) -> Dict[str, Any]:
        """EPUBファイルの目次構造を解析"""
        if not V2_SUPPORT:
            raise ImportError("v2.0 modules required for EPUB analysis")
            
        with open(epub_path, 'rb') as f:
            epub_analyzer = SplitEpubWordTOC(f)
            toc_map = epub_analyzer.get_enhanced_toc_map()
            
            # EPUB目次をTOCEntryに変換
            self.toc_entries = []
            self.level_stats = defaultdict(int)
            
            if epub_analyzer.toc_processor:
                for level in range(1, 4):
                    for entry in epub_analyzer.toc_processor.word_toc_generator.levels[level]:
                        toc_entry = TOCEntry(
                            text=entry['text'],
                            level=level,
                            file_href=entry.get('href', ''),
                            anchor=entry.get('anchor', ''),
                            parent_path=entry.get('hierarchy_path', '')
                        )
                        self.toc_entries.append(toc_entry)
                        self.level_stats[level] += 1
        
        return self._generate_analysis_report()
    
    def _detect_heading_styles(self, document: DocxDocument) -> Dict[str, int]:
        """見出しスタイルを検出してレベルにマッピング"""
        heading_styles = {}
        
        # 標準的な見出しスタイル
        standard_headings = {
            'Heading 1': 1, 'Heading 2': 2, 'Heading 3': 3,
            'Heading 4': 4, 'Heading 5': 5, 'Heading 6': 6,
            '見出し 1': 1, '見出し 2': 2, '見出し 3': 3,
            'Title': 1, 'Subtitle': 2
        }
        
        # 使用されているスタイルを確認
        used_styles = set()
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                used_styles.add(paragraph.style.name)
        
        # 標準スタイルのマッピング
        for style_name, level in standard_headings.items():
            if style_name in used_styles:
                heading_styles[style_name] = level
        
        # カスタムスタイルの検出（ヒューリスティック）
        for style_name in used_styles:
            if style_name not in heading_styles:
                # スタイル名パターンで判定
                if re.search(r'(chapter|章|section|節)', style_name, re.IGNORECASE):
                    heading_styles[style_name] = 1
                elif re.search(r'(subsection|小節|部)', style_name, re.IGNORECASE):
                    heading_styles[style_name] = 2
        
        return heading_styles
    
    def _build_hierarchy(self):
        """階層構造を構築"""
        if not self.toc_entries:
            return
            
        # 親子関係の構築
        stack = []  # (level, entry) のスタック
        
        for entry in self.toc_entries:
            # 現在のレベル以上の要素をスタックから除去
            while stack and stack[-1][0] >= entry.level:
                stack.pop()
            
            # 親パスの構築
            if stack:
                parent_entry = stack[-1][1]
                entry.parent_path = parent_entry.full_path
                entry.full_path = f"{entry.parent_path}/{entry.text}"
            
            stack.append((entry.level, entry))
        
        # 統計更新
        self.max_depth = max(self.level_stats.keys()) if self.level_stats else 0
        
        # 推奨分割レベルの決定
        self._calculate_recommended_split_level()
    
    def _calculate_recommended_split_level(self):
        """推奨分割レベルを計算"""
        if not self.level_stats:
            self.recommended_split_level = 1
            return
        
        # レベル1の数が適度（3-20章）なら1で分割
        level1_count = self.level_stats.get(1, 0)
        if 3 <= level1_count <= 20:
            self.recommended_split_level = 1
        # レベル1が少なくてレベル2が多いなら2で分割
        elif level1_count < 3 and self.level_stats.get(2, 0) > 5:
            self.recommended_split_level = 2
        # レベル1が多すぎるなら2で分割を提案
        elif level1_count > 20:
            self.recommended_split_level = 2
        else:
            self.recommended_split_level = 1
    
    def _generate_analysis_report(self) -> Dict[str, Any]:
        """分析レポートを生成"""
        report = {
            'total_entries': len(self.toc_entries),
            'max_depth': self.max_depth,
            'level_stats': dict(self.level_stats),
            'recommended_split_level': self.recommended_split_level,
            'entries': [
                {
                    'text': entry.text,
                    'level': entry.level,
                    'full_path': entry.full_path
                } for entry in self.toc_entries
            ],
            'split_preview': self._generate_split_preview()
        }
        return report
    
    def _generate_split_preview(self) -> List[Dict]:
        """分割プレビューを生成"""
        preview = []
        split_level = self.recommended_split_level
        
        current_section = None
        subsections = []
        
        for entry in self.toc_entries:
            if entry.level == split_level:
                # 前のセクションを保存
                if current_section:
                    preview.append({
                        'title': current_section.text,
                        'subsections': len(subsections),
                        'subsection_list': [s.text for s in subsections[:5]]  # 最初の5個のみ
                    })
                
                # 新しいセクション開始
                current_section = entry
                subsections = []
            elif entry.level > split_level and current_section:
                subsections.append(entry)
        
        # 最後のセクション
        if current_section:
            preview.append({
                'title': current_section.text,
                'subsections': len(subsections),
                'subsection_list': [s.text for s in subsections[:5]]
            })
        
        return preview

class WordDocumentSplitter:
    """Wordファイル分割クラス"""
    
    def __init__(self, analyzer: DynamicTOCAnalyzer):
        self.analyzer = analyzer
        self.original_document = None
        self.split_documents = []
        
    def split_document(self, docx_path: str, config: SplitConfig) -> List[str]:
        """Wordファイルを目次レベルに基づいて分割"""
        if not WORD_SUPPORT:
            raise ImportError("python-docx is required for Word splitting")
        
        self.original_document = Document(docx_path)
        output_files = []
        
        # 分割ポイントの決定
        split_points = self._determine_split_points(config.split_level)
        
        # 各セクションを分割
        for i, (start_idx, end_idx, title) in enumerate(split_points):
            section_doc = self._create_section_document(start_idx, end_idx, title, config)
            
            # ファイル名生成
            safe_title = self._sanitize_filename(title)
            filename = config.filename_pattern.format(
                index=i+1,
                title=safe_title[:config.max_filename_length]
            )
            
            output_path = Path(config.output_dir) / f"{filename}.docx"
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            section_doc.save(str(output_path))
            output_files.append(str(output_path))
        
        return output_files
    
    def _determine_split_points(self, split_level: int) -> List[Tuple[int, int, str]]:
        """分割ポイントを決定"""
        split_points = []
        current_start = 0
        
        for i, entry in enumerate(self.analyzer.toc_entries):
            if entry.level == split_level:
                # 前のセクションを終了
                if i > 0:  # 最初のエントリでない場合
                    prev_end = entry.paragraph_index
                    split_points.append((current_start, prev_end, prev_title))
                
                # 新しいセクション開始
                current_start = entry.paragraph_index
                prev_title = entry.text
        
        # 最後のセクション
        if self.analyzer.toc_entries:
            split_points.append((current_start, len(self.original_document.paragraphs), prev_title))
        
        return split_points
    
    def _create_section_document(self, start_idx: int, end_idx: int, title: str, config: SplitConfig) -> DocxDocument:
        """セクション文書を作成"""
        section_doc = Document()
        
        # タイトル追加
        title_para = section_doc.add_heading(title, level=1)
        
        # 元文書の段落をコピー
        for i in range(start_idx, min(end_idx, len(self.original_document.paragraphs))):
            original_para = self.original_document.paragraphs[i]
            
            if config.preserve_formatting:
                # 書式を保持してコピー
                new_para = section_doc.add_paragraph()
                new_para.style = original_para.style
                
                for run in original_para.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
            else:
                # テキストのみコピー
                section_doc.add_paragraph(original_para.text)
        
        return section_doc
    
    def _sanitize_filename(self, filename: str) -> str:
        """ファイル名の無効文字を除去"""
        # 無効文字を除去
        safe_name = re.sub(r'[<>:"/\\|?*]', '_', filename)
        safe_name = re.sub(r'\s+', '_', safe_name)
        safe_name = safe_name.strip('._')
        
        # 空の場合のデフォルト名
        if not safe_name:
            safe_name = "section"
        
        return safe_name

class MultiFormatExporter:
    """複数形式出力クラス"""
    
    def __init__(self, analyzer: DynamicTOCAnalyzer):
        self.analyzer = analyzer
    
    def export_to_pdf(self, content_sections: List[Dict], output_path: str) -> str:
        """PDF形式で出力"""
        if not PDF_SUPPORT:
            raise ImportError("reportlab is required for PDF export")
        
        doc = SimpleDocTemplate(output_path, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # 日本語フォント設定（オプション）
        try:
            # Windows環境での日本語フォント
            font_paths = [
                "C:/Windows/Fonts/msgothic.ttc",  # MS ゴシック
                "C:/Windows/Fonts/meiryo.ttc",    # メイリオ
            ]
            for font_path in font_paths:
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont('Japanese', font_path))
                    break
        except:
            pass  # フォント登録に失敗しても続行
        
        for section in content_sections:
            # セクションタイトル
            title = RLParagraph(section['title'], styles['Title'])
            story.append(title)
            story.append(Spacer(1, 12))
            
            # セクション内容
            for paragraph_text in section.get('content', []):
                para = RLParagraph(paragraph_text, styles['Normal'])
                story.append(para)
                story.append(Spacer(1, 6))
            
            story.append(PageBreak())
        
        doc.build(story)
        return output_path
    
    def export_to_epub(self, content_sections: List[Dict], output_path: str, metadata: Dict = None) -> str:
        """EPUB形式で出力"""
        if not EPUB_SUPPORT:
            raise ImportError("ebooklib is required for EPUB export")
        
        book = epub.EpubBook()
        
        # メタデータ設定
        metadata = metadata or {}
        book.set_identifier(f"epub_split_{datetime.now().strftime('%Y%m%d_%H%M%S')}")
        book.set_title(metadata.get('title', 'Split Document'))
        book.set_language(metadata.get('language', 'ja'))
        book.add_author(metadata.get('author', 'Unknown'))
        
        # 章を追加
        spine = ['nav']
        toc = []
        
        for i, section in enumerate(content_sections):
            # 章の作成
            chapter_filename = f"chapter_{i+1}.xhtml"
            chapter = epub.EpubHtml(
                title=section['title'],
                file_name=chapter_filename,
                lang='ja'
            )
            
            # HTML内容生成
            content_html = f"<h1>{section['title']}</h1>"
            for paragraph_text in section.get('content', []):
                content_html += f"<p>{paragraph_text}</p>"
            
            chapter.set_content(content_html)
            book.add_item(chapter)
            spine.append(chapter)
            toc.append(epub.Link(chapter_filename, section['title'], f"chapter_{i+1}"))
        
        # 目次設定
        book.toc = toc
        book.spine = spine
        
        # ナビゲーションファイル追加
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # EPUB書き出し
        epub.write_epub(output_path, book)
        return output_path

def main():
    """v3.0テスト実行"""
    print("📚 EPUB目次解析ツール v3.0 - 分割機能テスト")
    print("=" * 60)
    
    # 動的TOC解析のテスト
    analyzer = DynamicTOCAnalyzer()
    
    # 使用例の表示
    print("💡 使用例:")
    print("1. Wordファイル分析:")
    print("   analyzer = DynamicTOCAnalyzer()")
    print("   report = analyzer.analyze_word_document('document.docx')")
    print("")
    print("2. 分割実行:")
    print("   splitter = WordDocumentSplitter(analyzer)")
    print("   config = SplitConfig(split_level=1, output_format='word')")
    print("   files = splitter.split_document('document.docx', config)")
    print("")
    print("3. 複数形式出力:")
    print("   exporter = MultiFormatExporter(analyzer)")
    print("   exporter.export_to_pdf(sections, 'output.pdf')")
    print("   exporter.export_to_epub(sections, 'output.epub')")
    
    return True

if __name__ == "__main__":
    main()
